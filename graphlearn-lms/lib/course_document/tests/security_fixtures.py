"""Safe synthetic fixtures only; contains no real malware."""

import io
import zipfile
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject


def pdf(*, encrypted: bool = False, active: bool = False, pages: int = 1) -> bytes:
    page_objects = b"\n".join(b"<< /Type /Page >>" for _ in range(pages))
    markers = (b" /Encrypt <<>>" if encrypted else b"") + (b" /JavaScript /JS" if active else b"")
    return b"%PDF-1.7\n" + page_objects + markers + b"\n" + (b"safe-course-content\n" * 8) + b"%%EOF"


def docx(*, traversal: bool = False, macro: bool = False, xxe: bool = False, ordinary_zip: bool = False) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as package:
        if ordinary_zip:
            package.writestr("notes.txt", "ordinary zip")
        else:
            xml = '<!DOCTYPE x [<!ENTITY e SYSTEM "http://127.0.0.1/">]><document>&e;</document>' if xxe else "<document>safe course content</document>"
            package.writestr("[Content_Types].xml", "<Types></Types>")
            package.writestr("_rels/.rels", "<Relationships></Relationships>")
            package.writestr("word/document.xml", xml)
            if traversal:
                package.writestr("../outside.txt", "blocked")
            if macro:
                package.writestr("word/vbaProject.bin", b"safe synthetic macro marker")
    return output.getvalue()


def text_pdf(pages: list[str]) -> bytes:
    output = io.BytesIO()
    writer = PdfWriter()
    for text in pages:
        page = writer.add_blank_page(width=612, height=792)
        font = DictionaryObject({
            NameObject("/Type"): NameObject("/Font"), NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        })
        resources = DictionaryObject({NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)})})
        page[NameObject("/Resources")] = resources
        escaped = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        stream = DecodedStreamObject()
        stream.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1", "replace"))
        page[NameObject("/Contents")] = writer._add_object(stream)
    writer.write(output)
    return output.getvalue()


def blank_pdf(pages: int = 1) -> bytes:
    output = io.BytesIO(); writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=612, height=792)
    writer.write(output); return output.getvalue()


def structured_docx() -> bytes:
    output = io.BytesIO(); document = Document()
    document.add_heading("Database Normalization", level=1)
    document.add_paragraph("First paragraph with ﬁ ligature and   extra spaces.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Form"; table.cell(0, 1).text = "Meaning"
    table.cell(1, 0).text = "1NF"; table.cell(1, 1).text = "Atomic values"
    document.add_paragraph("Final paragraph after the table.")
    document.add_paragraph("ACID", style="List Bullet")
    document.save(output); return output.getvalue()
