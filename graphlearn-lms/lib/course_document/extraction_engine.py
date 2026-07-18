"""Network-free parser adapters and deterministic text normalization."""

from __future__ import annotations

import io
import json
import re
import unicodedata
from typing import Any

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
import pypdf
import docx


_CONTROL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_SPACES = re.compile(r"[^\S\n]+")


def normalize_text(value: str) -> tuple[str, int]:
    value = unicodedata.normalize("NFC", value).replace("\r\n", "\n").replace("\r", "\n")
    value = value.replace("\u200b", "").replace("\ufeff", "").replace("ﬁ", "fi").replace("ﬂ", "fl")
    removed = len(_CONTROL.findall(value))
    value = _CONTROL.sub("", value)
    lines = [_SPACES.sub(" ", line).strip() for line in value.split("\n")]
    # Conservative line-break dehyphenation only for lowercase continuation.
    joined: list[str] = []
    for line in lines:
        if joined and joined[-1].endswith("-") and line[:1].islower():
            joined[-1] = joined[-1][:-1] + line
        elif line:
            joined.append(line)
    return "\n".join(joined).strip(), removed


def pdf_version() -> str:
    return pypdf.__version__


def docx_version() -> str:
    return docx.__version__


def extract_pdf(payload: bytes) -> dict[str, Any]:
    reader = PdfReader(io.BytesIO(payload), strict=True)
    blocks: list[dict[str, Any]] = []
    warnings: list[str] = []
    pages_with_text = pages_without_text = pages_with_images = 0
    page_texts: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        raw = "" if page.get_contents() is None else (page.extract_text(extraction_mode="layout") or "")
        text, removed = normalize_text(raw)
        if removed > 20:
            warnings.append("CONTROL_CHARACTERS_REMOVED")
        image_count = len(getattr(page, "images", []))
        pages_with_images += int(image_count > 0)
        if text:
            pages_with_text += 1
            page_texts.append(text)
            for block_number, paragraph in enumerate(re.split(r"\n\s*\n|\n", text), start=1):
                normalized, _ = normalize_text(paragraph)
                if normalized:
                    blocks.append({
                        "block_type": "PARAGRAPH", "text": normalized,
                        "page_start": page_number, "page_end": page_number,
                        "source_locator": f"pdf:page={page_number}:block={block_number}",
                        "metadata_json": json.dumps({"image_count": image_count}, sort_keys=True),
                    })
        else:
            pages_without_text += 1
    total_pages = len(reader.pages)
    if pages_without_text and pages_with_text:
        warnings.append("PDF_PARTIAL_TEXT_LAYER")
    if total_pages and sum(len(text) for text in page_texts) / total_pages < 40:
        warnings.append("PDF_LOW_TEXT_DENSITY")
    return {
        "blocks": blocks, "warnings": sorted(set(warnings)), "total_pages": total_pages,
        "pages_with_text": pages_with_text, "pages_without_text": pages_without_text,
        "pages_with_images": pages_with_images,
        "extractor_name": "pdf:pypdf", "extractor_version": pypdf.__version__,
        "parser_metadata_json": json.dumps({"engine": "pypdf"}, sort_keys=True),
    }


def _paragraph_block(paragraph: Paragraph, index: int) -> dict[str, Any] | None:
    text, _ = normalize_text(paragraph.text)
    if not text:
        return None
    style = paragraph.style.name if paragraph.style is not None else ""
    heading = re.match(r"Heading\s+(\d+)", style, re.IGNORECASE)
    block_type = "HEADING" if heading else ("LIST_ITEM" if style.lower().startswith("list") else ("CAPTION" if "caption" in style.lower() else "PARAGRAPH"))
    level = int(heading.group(1)) if heading else None
    return {
        "block_type": block_type, "text": text, "heading_level": level,
        "heading_text": text if heading else None,
        "source_locator": f"docx:body=paragraph:{index}",
        "metadata_json": json.dumps({"style": style}, sort_keys=True),
    }


def _table_block(table: Table, index: int) -> dict[str, Any] | None:
    rows: list[str] = []
    max_columns = 0
    for row in table.rows:
        cells = [normalize_text(cell.text)[0] for cell in row.cells]
        max_columns = max(max_columns, len(cells))
        if any(cells):
            rows.append(" | ".join(cells))
    if not rows:
        return None
    return {
        "block_type": "TABLE_TEXT", "text": "\n".join(rows),
        "source_locator": f"docx:body=table:{index}",
        "metadata_json": json.dumps({"rows": len(rows), "columns": max_columns}, sort_keys=True),
    }


def extract_docx(payload: bytes) -> dict[str, Any]:
    document = Document(io.BytesIO(payload))
    blocks: list[dict[str, Any]] = []
    warnings: list[str] = []
    paragraph_index = table_index = 0
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph_index += 1
            block = _paragraph_block(Paragraph(child, document), paragraph_index)
        elif child.tag.endswith("}tbl"):
            table_index += 1
            block = _table_block(Table(child, document), table_index)
        else:
            block = None
        if block is not None:
            blocks.append(block)
    for section_index, section in enumerate(document.sections, start=1):
        for kind, paragraphs in (("HEADER", section.header.paragraphs), ("FOOTER", section.footer.paragraphs)):
            for index, paragraph in enumerate(paragraphs, start=1):
                text, _ = normalize_text(paragraph.text)
                if text:
                    blocks.append({"block_type": kind, "text": text, "source_locator": f"docx:section={section_index}:{kind.lower()}:{index}"})
    xml = document.element.xml
    if "<w:ins" in xml or "<w:del" in xml:
        warnings.append("DOCX_TRACKED_CHANGES_PRESENT")
    return {
        "blocks": blocks, "warnings": warnings, "total_pages": None,
        "pages_with_text": 0, "pages_without_text": 0, "pages_with_images": 0,
        "extractor_name": "docx:python-docx", "extractor_version": docx.__version__,
        "parser_metadata_json": json.dumps({"engine": "python-docx"}, sort_keys=True),
    }
